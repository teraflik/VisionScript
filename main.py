import io
import os

# Imports the Google Cloud client library
from google.cloud import vision
from google.cloud.vision import types

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def VisionScript(indir):
	# Instantiates a client
	client = vision.ImageAnnotatorClient()

	folders = []
	files = []
	
	print("\nReading Directory images\\ ... ", end="")
	for entry in os.scandir(indir):
		if entry.is_dir():
			folders.append(entry.path)
		elif entry.is_file():
			files.append(entry.path)
	print("done\n")
	
	i = 0
	total = len(files)
	#output_txt = ""
	#f = open('output_old.txt', 'w', encoding='utf-8')

	#Preparing docx file with formatting and margins.
	document = Document()
	section = document.sections[0]
	section.left_margin, section.right_margin = (Cm(1.27), Cm(1.27))
	section.top_margin, section.bottom_margin = (Cm(1.27), Cm(1.27))
	section.gutter = 0

	style = document.styles['Normal']
	font = style.font
	font.name = 'Kokila'
	font.size = Pt(18)
	
	document2 = Document()

	style2 = document2.styles['Normal']
	font = style2.font
	font.name = 'Kokila'
	font.size = Pt(18)

	for image in files:
		i = i+1
		file_name = str(image)
		print("Processing {:d} of {:d} images: {:s}...".format(i, total, file_name), end="")

		with io.open(file_name, 'rb') as image_file:
			content = image_file.read()

		image = types.Image(content=content)
		
		response = client.document_text_detection(image=image)
		original_output = response.full_text_annotation.text
		
		#Replace newline characters with space
		new_output = original_output.replace("\r", "")
		new_output = new_output.replace("\n", " ")

		#generate output for both type of files
		#output_txt = original_output + "\n\n"
		#f.write(output_txt)

		paragraph = document.add_paragraph(new_output + "\n\n")
		paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
		paragraph.style = document.styles['Normal']

		paragraph2 = document2.add_paragraph(original_output + "\n\n")
		paragraph2.style = document2.styles['Normal']
		
		print("... done!")
	
	document.save('Word.docx')
	document2.save('Word_old.docx')
	#f.close()
	print("\nAll images processed. Output saved to Word.docx, Word_old.docx\n")

if __name__ == "__main__":
	VisionScript('images/')
	input("Press Enter to continue...\n") 